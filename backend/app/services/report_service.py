"""报告生成服务：偏离报告 / 风险报告 Docx 导出

生成含证据引用 + 原文片段 + 统计摘要的 Word 文档，上传 MinIO 返回 file_key。
"""
import io
import json
import uuid
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.minio_client import minio_client
from app.models.comparison import ComparisonResult, ComparisonTask, ComparisonVerdict
from app.models.contract import Contract, ContractRisk, RiskLevel

# 判定结果中文标签
VERDICT_LABELS = {
    ComparisonVerdict.MATCH: "一致",
    ComparisonVerdict.DEVIATION: "偏离",
    ComparisonVerdict.MISSING: "缺失",
    ComparisonVerdict.EXTRA: "多余",
    ComparisonVerdict.NEED_CONFIRM: "待确认",
}

# 风险等级中文标签
RISK_LEVEL_LABELS = {
    RiskLevel.HIGH: "高",
    RiskLevel.MEDIUM: "中",
    RiskLevel.LOW: "低",
    RiskLevel.INFO: "提示",
}

# Docx MIME 类型
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# 高风险 / 偏离 红色
RED = RGBColor(0xFF, 0x00, 0x00)

# 原文片段最大长度
SNIPPET_MAX_LEN = 200


class ReportService:
    """生成偏离报告和风险报告 Docx 并上传 MinIO"""

    # ------------------------------------------------------------------
    # 公共接口
    # ------------------------------------------------------------------
    async def generate_comparison_report(self, session: AsyncSession, task_id: uuid.UUID) -> str:
        """生成参数偏离比对报告 Docx，上传 MinIO，返回 file_key

        - 标题、任务信息、统计摘要（一致/偏离/缺失数量）
        - 每条参数：参数名、招标要求、规格响应、判定结果、置信度、理由、原文引用
        """
        task = await session.get(ComparisonTask, task_id)
        if not task:
            raise ValueError(f"比对任务不存在: {task_id}")

        result = await session.execute(
            select(ComparisonResult)
            .where(ComparisonResult.task_id == task_id)
            .order_by(ComparisonResult.created_at)
        )
        results: List[ComparisonResult] = list(result.scalars().all())

        doc = self._build_comparison_doc(task, results)
        file_key = self._make_key("comparison", task_id)
        self._upload_docx(doc, file_key)

        # 回写 report_key 便于后续直接下载
        task.report_key = file_key
        await session.commit()
        return file_key

    async def generate_risk_report(self, session: AsyncSession, contract_id: uuid.UUID) -> str:
        """生成合同风险扫描报告 Docx，上传 MinIO，返回 file_key

        - 标题、合同信息、统计摘要（高/中/低数量）
        - 每条风险：类别、等级、标题、描述、修改建议、置信度、原文引用
        """
        contract = await session.get(Contract, contract_id)
        if not contract:
            raise ValueError(f"合同不存在: {contract_id}")

        result = await session.execute(
            select(ContractRisk)
            .where(ContractRisk.contract_id == contract_id)
            .order_by(ContractRisk.created_at)
        )
        risks: List[ContractRisk] = list(result.scalars().all())

        doc = self._build_risk_doc(contract, risks)
        file_key = self._make_key("risk", contract_id)
        self._upload_docx(doc, file_key)
        return file_key

    def presigned_url(self, file_key: str, expires_days: int = 7) -> Optional[str]:
        """生成 MinIO 预签名下载 URL（失败返回 None）"""
        try:
            return minio_client.presigned_get_object(
                settings.MINIO_BUCKET, file_key, expires=timedelta(days=expires_days)
            )
        except Exception:
            return None

    # ------------------------------------------------------------------
    # 偏离报告构建
    # ------------------------------------------------------------------
    def _build_comparison_doc(self, task: ComparisonTask, results: List[ComparisonResult]) -> Document:
        doc = Document()

        # 标题
        title = doc.add_heading("参数偏离比对报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 一、任务信息
        doc.add_heading("一、任务信息", level=1)
        info_lines: List[tuple] = [
            ("任务ID", str(task.id)),
            ("项目ID", str(task.project_id) if task.project_id else "-"),
            ("招标文件ID", str(task.tender_doc_id) if task.tender_doc_id else "-"),
        ]
        if task.spec_doc_id:
            info_lines.append(("规格书ID", str(task.spec_doc_id)))
        if task.product_id:
            info_lines.append(("产品ID", str(task.product_id)))
        info_lines.append(("任务状态", task.status))
        info_lines.append(("生成时间", _now_str()))
        for label, value in info_lines:
            self._add_field(doc, label, value)

        # 二、统计摘要
        doc.add_heading("二、统计摘要", level=1)
        summary = self._compute_comparison_summary(results)
        headers = ["总参数数", "一致", "偏离", "缺失", "多余", "待确认"]
        values = [summary["total"], summary["match"], summary["deviation"],
                  summary["missing"], summary["extra"], summary["need_confirm"]]
        self._add_summary_table(doc, headers, values)

        # 三、参数明细
        doc.add_heading("三、参数明细", level=1)
        if not results:
            doc.add_paragraph("（暂无比对结果）")
        else:
            # 概览表
            overview = doc.add_table(rows=1, cols=5)
            overview.style = "Table Grid"
            for i, h in enumerate(["序号", "参数名", "判定", "置信度", "招标要求/规格响应"]):
                overview.rows[0].cells[i].text = h
            for idx, r in enumerate(results, 1):
                cells = overview.add_row().cells
                cells[0].text = str(idx)
                cells[1].text = r.param_name or "-"
                cells[2].text = VERDICT_LABELS.get(r.verdict, str(r.verdict))
                cells[3].text = _fmt_confidence(r.confidence)
                brief = (r.tender_value or "-")[:30]
                cells[4].text = brief

            doc.add_paragraph()  # 空行分隔

            # 每条参数详情
            for idx, r in enumerate(results, 1):
                doc.add_heading(f"{idx}. {r.param_name or '(未命名参数)'}", level=2)
                self._add_field(doc, "招标要求", r.tender_value)
                self._add_field(doc, "规格响应", r.spec_value)

                # 判定结果（偏离/缺失标红）
                p = doc.add_paragraph()
                rr = p.add_run("判定结果：")
                rr.bold = True
                verdict_text = VERDICT_LABELS.get(r.verdict, str(r.verdict))
                vr = p.add_run(verdict_text)
                if r.verdict in (ComparisonVerdict.DEVIATION, ComparisonVerdict.MISSING):
                    vr.font.color.rgb = RED

                # 置信度
                p = doc.add_paragraph()
                rr = p.add_run("置信度：")
                rr.bold = True
                p.add_run(_fmt_confidence(r.confidence))

                # 判断理由
                self._add_field(doc, "判断理由", r.reason)

                # 证据引用
                self._add_evidence_block(doc, "招标文件证据", r.tender_evidence)
                self._add_evidence_block(doc, "规格书证据", r.spec_evidence)

        return doc

    # ------------------------------------------------------------------
    # 风险报告构建
    # ------------------------------------------------------------------
    def _build_risk_doc(self, contract: Contract, risks: List[ContractRisk]) -> Document:
        doc = Document()

        # 标题
        title = doc.add_heading("合同风险扫描报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 一、合同信息
        doc.add_heading("一、合同信息", level=1)
        info_lines: List[tuple] = [
            ("合同ID", str(contract.id)),
            ("合同标题", contract.title or "-"),
            ("相对方", contract.counterparty or "-"),
        ]
        if contract.sign_date:
            info_lines.append(("签订日期", str(contract.sign_date)))
        if contract.effective_date:
            info_lines.append(("生效日期", str(contract.effective_date)))
        if contract.expire_date:
            info_lines.append(("到期日期", str(contract.expire_date)))
        if contract.amount is not None:
            info_lines.append(("合同金额", str(contract.amount)))
        info_lines.append(("生成时间", _now_str()))
        for label, value in info_lines:
            self._add_field(doc, label, value)

        # 二、统计摘要
        doc.add_heading("二、统计摘要", level=1)
        summary = self._compute_risk_summary(risks)
        headers = ["风险总数", "高风险", "中风险", "低风险", "提示"]
        values = [summary["total"], summary["high"], summary["medium"],
                  summary["low"], summary["info"]]
        self._add_summary_table(doc, headers, values)

        # 三、风险明细
        doc.add_heading("三、风险明细", level=1)
        if not risks:
            doc.add_paragraph("（暂无风险扫描结果）")
        else:
            # 概览表
            overview = doc.add_table(rows=1, cols=5)
            overview.style = "Table Grid"
            for i, h in enumerate(["序号", "类别", "等级", "标题", "置信度"]):
                overview.rows[0].cells[i].text = h
            for idx, r in enumerate(risks, 1):
                cells = overview.add_row().cells
                cells[0].text = str(idx)
                cells[1].text = r.category or "-"
                cells[2].text = RISK_LEVEL_LABELS.get(r.level, str(r.level))
                cells[3].text = r.title or "-"
                cells[4].text = _fmt_confidence(r.confidence)

            doc.add_paragraph()  # 空行分隔

            # 每条风险详情
            for idx, r in enumerate(risks, 1):
                heading_text = (
                    f"{idx}. [{RISK_LEVEL_LABELS.get(r.level, str(r.level))}] "
                    f"{r.title or '(未命名风险)'}"
                )
                h = doc.add_heading(heading_text, level=2)
                # 高风险标题红色
                if r.level == RiskLevel.HIGH:
                    for run in h.runs:
                        run.font.color.rgb = RED

                self._add_field(doc, "风险类别", r.category)
                self._add_field(doc, "风险描述", r.description)
                self._add_field(doc, "修改建议", r.suggestion)

                # 置信度
                p = doc.add_paragraph()
                rr = p.add_run("置信度：")
                rr.bold = True
                p.add_run(_fmt_confidence(r.confidence))

                # 证据引用
                self._add_evidence_block(doc, "原文证据", r.evidence)

        return doc

    # ------------------------------------------------------------------
    # 辅助方法
    # ------------------------------------------------------------------
    def _add_field(self, doc: Document, label: str, value: Optional[str]):
        """添加一个 字段名：值 段落"""
        p = doc.add_paragraph()
        r = p.add_run(f"{label}：")
        r.bold = True
        p.add_run(value if value else "-")

    def _add_summary_table(self, doc: Document, headers: List[str], values: List[Any]):
        """添加统计摘要表"""
        table = doc.add_table(rows=2, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for i, v in enumerate(values):
            table.rows[1].cells[i].text = str(v)

    def _add_evidence_block(self, doc: Document, label: str, evidence: Any):
        """添加证据引用区块（标题 + 格式化原文片段）"""
        formatted = self._format_evidence(evidence)
        p = doc.add_paragraph()
        r = p.add_run(f"{label}：")
        r.bold = True
        if formatted:
            p.add_run(formatted)
        else:
            p.add_run("无")

    def _format_evidence(self, evidence: Any) -> str:
        """格式化证据引用列表为可读文本

        兼容字段名：
        - 原文片段：content / snippet
        - 页码：page_number / page
        """
        if not evidence:
            return ""
        # 统一为列表
        if isinstance(evidence, dict):
            items = [evidence]
        elif isinstance(evidence, list):
            items = evidence
        else:
            return str(evidence)

        lines: List[str] = []
        for i, ev in enumerate(items, 1):
            if not isinstance(ev, dict):
                lines.append(f"[{i}] {ev}")
                continue

            content = ev.get("content") or ev.get("snippet") or ""
            page = ev.get("page_number") or ev.get("page")
            section = ev.get("section")
            chunk_id = ev.get("chunk_id")
            document_id = ev.get("document_id")
            table_ref = ev.get("table_ref")

            loc_parts: List[str] = []
            if page:
                loc_parts.append(f"第{page}页")
            if section:
                loc_parts.append(f"章节:{section}")
            if table_ref:
                loc_parts.append(table_ref)
            if chunk_id:
                loc_parts.append(f"chunk:{chunk_id}")
            if document_id:
                loc_parts.append(f"doc:{document_id}")
            loc = " | ".join(loc_parts)

            if content:
                snippet = content if len(content) <= SNIPPET_MAX_LEN else content[:SNIPPET_MAX_LEN] + "…"
                lines.append(f"[{i}] {loc}\n原文片段：{snippet}")
            else:
                lines.append(f"[{i}] {loc}")
        return "\n".join(lines)

    def _compute_comparison_summary(self, results: List[ComparisonResult]) -> dict:
        """计算比对统计摘要（一致/偏离/缺失/多余/待确认）"""
        summary = {
            "total": len(results),
            "match": 0,
            "deviation": 0,
            "missing": 0,
            "extra": 0,
            "need_confirm": 0,
        }
        for r in results:
            if r.verdict == ComparisonVerdict.MATCH:
                summary["match"] += 1
            elif r.verdict == ComparisonVerdict.DEVIATION:
                summary["deviation"] += 1
            elif r.verdict == ComparisonVerdict.MISSING:
                summary["missing"] += 1
            elif r.verdict == ComparisonVerdict.EXTRA:
                summary["extra"] += 1
            elif r.verdict == ComparisonVerdict.NEED_CONFIRM:
                summary["need_confirm"] += 1
        return summary

    def _compute_risk_summary(self, risks: List[ContractRisk]) -> dict:
        """计算风险统计摘要（高/中/低/提示）"""
        summary = {"total": len(risks), "high": 0, "medium": 0, "low": 0, "info": 0}
        for r in risks:
            if r.level == RiskLevel.HIGH:
                summary["high"] += 1
            elif r.level == RiskLevel.MEDIUM:
                summary["medium"] += 1
            elif r.level == RiskLevel.LOW:
                summary["low"] += 1
            elif r.level == RiskLevel.INFO:
                summary["info"] += 1
        return summary

    def _make_key(self, kind: str, entity_id: uuid.UUID) -> str:
        """生成 MinIO file_key"""
        ts = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        return f"reports/{kind}/{entity_id}/{ts}.docx"

    def _upload_docx(self, doc: Document, file_key: str):
        """将 docx 保存到内存并上传 MinIO"""
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        minio_client.put_object(
            settings.MINIO_BUCKET,
            file_key,
            io.BytesIO(data),
            length=len(data),
            content_type=DOCX_CONTENT_TYPE,
        )

    # ------------------------------------------------------------------
    # 多产品对比报告（横向 A4 布局）
    # ------------------------------------------------------------------
    async def generate_multi_comparison_report(
        self,
        session: AsyncSession,
        task_ids: List[uuid.UUID],
        include_ai_summary: bool = True,
    ) -> bytes:
        """生成多产品对比报告（横向 A4 布局）

        特性：
        - A4 横向布局（page_width=29.7cm, page_height=21cm, orientation=LANDSCAPE）
        - 概览表 + AI 智能选型总结（LLM 综合全表生成推荐选型理由）
        - 分产品独立矩阵
        - 证据引用块显示 source_file_id/evidence_span_id/page
        - 颜色映射（绿/黄/红）+ 标签映射（✔满足/☐部分满足/×不满足）
        """
        from docx import Document as DocxDocument
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from app.models.comparison import ComparisonTask, ComparisonResult, ComparisonVerdict
        from app.ai.llm.client import get_llm_client
        from app.ai.llm.response_parser import parse_llm_json, get_force_chinese_system_prompt

        # 加载所有任务的比对结果
        all_results: Dict[uuid.UUID, List[ComparisonResult]] = {}
        tasks_info: Dict[uuid.UUID, ComparisonTask] = {}
        for tid in task_ids:
            task = await session.get(ComparisonTask, tid)
            if not task:
                continue
            tasks_info[tid] = task
            stmt = select(ComparisonResult).where(ComparisonResult.task_id == tid)
            result = await session.execute(stmt)
            all_results[tid] = list(result.scalars().all())

        if not all_results:
            raise ValueError("无比对结果可生成报告")

        # 创建 Word 文档
        doc = DocxDocument()

        # 设置 A4 横向布局
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # 标题
        title = doc.add_heading("多产品参数对比报告", level=0)
        title.alignment = 1  # 居中

        # 概览表
        doc.add_heading("一、对比概览", level=1)
        overview_table = doc.add_table(rows=1, cols=len(task_ids) + 2)
        overview_table.style = "Light Grid Accent 1"
        overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = overview_table.rows[0].cells
        hdr[0].text = "产品/任务"
        hdr[1].text = "总参数数"
        for i, tid in enumerate(task_ids):
            hdr[i + 2].text = f"产品{i+1}"
        for tid, results in all_results.items():
            row = overview_table.add_row().cells
            row[0].text = str(tasks_info[tid].id)[:8]
            row[1].text = str(len(results))
            summary = tasks_info[tid].summary or {}
            match_count = summary.get("match", 0)
            deviation_count = summary.get("deviation", 0)
            missing_count = summary.get("missing", 0)
            row[2].text = f"满足:{match_count} 偏离:{deviation_count} 缺失:{missing_count}"

        # AI 智能选型总结
        if include_ai_summary:
            doc.add_heading("二、AI 智能选型总结", level=1)
            ai_summary = await self._generate_ai_selection_summary(all_results, tasks_info)
            doc.add_paragraph(ai_summary)

        # 分产品对比矩阵
        doc.add_heading("三、分产品对比明细", level=1)
        for idx, (tid, results) in enumerate(all_results.items(), 1):
            doc.add_heading(f"产品 {idx} 对比明细", level=2)
            await self._add_product_matrix(doc, results)
            doc.add_page_break()

        # 保存到字节
        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def _generate_ai_selection_summary(
        self,
        all_results: Dict[uuid.UUID, List],
        tasks_info: Dict[uuid.UUID, ComparisonTask],
    ) -> str:
        """LLM 综合全表生成推荐选型理由"""
        llm = get_llm_client()
        # 构建全表摘要
        summary_data = []
        for tid, results in all_results.items():
            task = tasks_info[tid]
            summary = task.summary or {}
            summary_data.append({
                "product_index": len(summary_data) + 1,
                "total": len(results),
                "match": summary.get("match", 0),
                "deviation": summary.get("deviation", 0),
                "missing": summary.get("missing", 0),
                "need_confirm": summary.get("need_confirm", 0),
                "satisfaction_rate": (
                    summary.get("match", 0) / len(results) if results else 0
                ),
            })
        prompt = (
            "你是产品选型专家。请基于以下多产品对比结果，给出推荐选型建议。\n\n"
            f"对比数据：{json.dumps(summary_data, ensure_ascii=False, indent=2)}\n\n"
            "请分析：\n"
            "1. 各产品的满足率排名\n"
            "2. 偏离和缺失项的影响评估\n"
            "3. 推荐选型方案及理由\n"
            "4. 风险提示\n\n"
            "输出 300-500 字的中文分析。"
        )
        messages = [
            {"role": "system", "content": "你是产品选型专家，擅长多产品对比分析。" + get_force_chinese_system_prompt()},
            {"role": "user", "content": prompt},
        ]
        try:
            return await llm.chat(messages, temperature=0.3, max_tokens=2048)
        except Exception as e:
            return f"AI 总结生成失败: {e}"

    async def _add_product_matrix(self, doc, results: List[ComparisonResult]):
        """添加单个产品的对比矩阵"""
        if not results:
            doc.add_paragraph("（无对比结果）")
            return

        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        headers = ["参数名", "招标要求", "规格响应", "判定", "置信度", "证据引用"]
        for i, h in enumerate(headers):
            hdr[i].text = h

        for r in results:
            row = table.add_row().cells
            row[0].text = r.param_name or "-"
            row[1].text = (r.tender_value or "-")[:100]
            row[2].text = (r.spec_value or "-")[:100]
            # 判定 + 颜色
            verdict_cell = row[3]
            verdict_text, color = self._verdict_label_and_color(r.verdict)
            verdict_cell.text = verdict_text
            if color:
                for paragraph in verdict_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = color
            row[4].text = f"{r.confidence:.0%}" if r.confidence else "-"
            # 证据引用
            evidence_text = self._format_evidence_refs(r.tender_evidence, r.spec_evidence, r.evidence_span_id)
            row[5].text = evidence_text

    @staticmethod
    def _verdict_label_and_color(verdict):
        """判定结果 → 标签 + 颜色"""
        from docx.shared import RGBColor
        mapping = {
            ComparisonVerdict.MATCH: ("✔ 满足", RGBColor(0, 128, 0)),       # 绿
            ComparisonVerdict.DEVIATION: ("× 偏离", RGBColor(255, 0, 0)),    # 红
            ComparisonVerdict.MISSING: ("× 缺失", RGBColor(255, 0, 0)),      # 红
            ComparisonVerdict.EXTRA: ("☐ 多余", RGBColor(255, 165, 0)),      # 橙
            ComparisonVerdict.NEED_CONFIRM: ("☐ 待确认", RGBColor(255, 165, 0)),  # 橙
        }
        return mapping.get(verdict, ("未知", None))

    @staticmethod
    def _format_evidence_refs(tender_ev, spec_ev, evidence_span_id=None) -> str:
        """格式化证据引用块"""
        refs = []
        if evidence_span_id:
            refs.append(f"[证据ID:{str(evidence_span_id)[:8]}]")
        if tender_ev and isinstance(tender_ev, list):
            for ev in tender_ev[:2]:
                if isinstance(ev, dict):
                    page = ev.get("page_number", "?")
                    refs.append(f"[招标 第{page}页]")
        if spec_ev and isinstance(spec_ev, list):
            for ev in spec_ev[:2]:
                if isinstance(ev, dict):
                    page = ev.get("page_number", "?")
                    refs.append(f"[规格 第{page}页]")
        return " ".join(refs) if refs else "-"


def _now_str() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")


def _fmt_confidence(confidence: Optional[float]) -> str:
    if confidence is None:
        return "-"
    return f"{confidence:.1%}"


_report_service: ReportService | None = None


def get_report_service() -> ReportService:
    global _report_service
    if _report_service is None:
        _report_service = ReportService()
    return _report_service
