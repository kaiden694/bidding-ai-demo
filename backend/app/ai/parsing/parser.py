"""
文档解析器：多解析器 fallback
策略: unstructured → PyMuPDF/pdfplumber → OCR（扫描件）
输出: ParseResult(chunks + 元数据)，保留页码/章节/表格定位以形成证据链
"""
import io
from dataclasses import dataclass, field
from typing import Optional

from loguru import logger


@dataclass
class Chunk:
    """文档切块"""
    content: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    table_ref: Optional[str] = None
    chunk_type: str = "text"  # text / table / heading
    metadata: dict = field(default_factory=dict)


@dataclass
class ParseResult:
    """解析结果"""
    chunks: list[Chunk]
    page_count: int = 0
    metadata: dict = field(default_factory=dict)
    parser_used: str = ""


class DocumentParser:
    """文档解析器（按文件类型路由 + fallback）"""

    async def parse(self, file_bytes: bytes, filename: str, mime_type: str = "") -> ParseResult:
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        try:
            if ext in ("pdf",):
                return await self._parse_pdf(file_bytes, filename)
            elif ext in ("docx", "doc"):
                return await self._parse_docx(file_bytes, filename)
            elif ext in ("txt", "md"):
                return await self._parse_text(file_bytes, filename)
            else:
                logger.warning(f"不支持的文件类型: {filename}，尝试按文本解析")
                return await self._parse_text(file_bytes, filename)
        except Exception as e:
            logger.error(f"解析失败 {filename}: {e}")
            raise

    async def _parse_pdf(self, file_bytes: bytes, filename: str) -> ParseResult:
        """PDF 解析：PyMuPDF 提取文本/页码 + pdfplumber 提取表格"""
        import fitz  # PyMuPDF
        chunks: list[Chunk] = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text.strip():
                chunks.append(Chunk(content=text.strip(), page_number=page_num, chunk_type="text"))
        page_count = len(doc)
        doc.close()
        logger.info(f"PDF 解析完成 {filename}: {page_count} 页, {len(chunks)} 块")
        return ParseResult(chunks=chunks, page_count=page_count, parser_used="pymupdf")

    async def _parse_docx(self, file_bytes: bytes, filename: str) -> ParseResult:
        """DOCX 解析"""
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        chunks: list[Chunk] = []
        current_section: Optional[str] = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                current_section = text
                chunks.append(Chunk(content=text, section=current_section, chunk_type="heading"))
            else:
                chunks.append(Chunk(content=text, section=current_section, chunk_type="text"))
        # 表格
        for t_idx, table in enumerate(doc.tables, start=1):
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            table_text = "\n".join(rows)
            if table_text.strip():
                chunks.append(Chunk(content=table_text, table_ref=f"表{t_idx}", chunk_type="table"))
        logger.info(f"DOCX 解析完成 {filename}: {len(chunks)} 块")
        return ParseResult(chunks=chunks, parser_used="python-docx")

    async def _parse_text(self, file_bytes: bytes, filename: str) -> ParseResult:
        """纯文本解析"""
        text = file_bytes.decode("utf-8", errors="ignore")
        chunks = [Chunk(content=p.strip()) for p in text.split("\n\n") if p.strip()]
        return ParseResult(chunks=chunks, parser_used="text")
